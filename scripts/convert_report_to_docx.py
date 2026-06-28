"""Convert PROJECT_REPORT.md to a formatted Word (.docx) document."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_heading_style(paragraph, level):
    colors = {
        1: RGBColor(0x0A, 0x6E, 0xBD),
        2: RGBColor(0x0A, 0x6E, 0xBD),
        3: RGBColor(0x1A, 0x1A, 0x2E),
        4: RGBColor(0x1A, 0x1A, 0x2E),
    }
    for run in paragraph.runs:
        run.font.color.rgb = colors.get(level, RGBColor(0, 0, 0))


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0A6EBD')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p


def set_table_style(table):
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)
            if i == 0:
                cell._tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), '0A6EBD')
                cell._tc.tcPr.append(shd)
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        run.font.bold = True


def parse_inline(run_or_text):
    """Strip inline markdown markers and return (text, bold, italic, code)."""
    text = run_or_text
    bold = False
    italic = False
    code = False
    if text.startswith('`') and text.endswith('`') and len(text) > 1:
        code = True
        text = text[1:-1]
    elif text.startswith('**') and text.endswith('**') and len(text) > 3:
        bold = True
        text = text[2:-2]
    elif text.startswith('*') and text.endswith('*') and len(text) > 1:
        italic = True
        text = text[1:-1]
    return text, bold, italic, code


def add_paragraph_with_inline(doc, raw_text, style=None, base_size=Pt(10.5)):
    """Add a paragraph parsing inline bold/italic/code markers."""
    para = doc.add_paragraph()
    if style:
        para.style = style
    # Split on inline code, bold, italic patterns
    tokens = re.split(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)', raw_text)
    for token in tokens:
        if not token:
            continue
        if token.startswith('`') and token.endswith('`'):
            run = para.add_run(token[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif token.startswith('**') and token.endswith('**'):
            run = para.add_run(token[2:-2])
            run.bold = True
            run.font.size = base_size
        elif token.startswith('*') and token.endswith('*'):
            run = para.add_run(token[1:-1])
            run.italic = True
            run.font.size = base_size
        else:
            run = para.add_run(token)
            run.font.size = base_size
    return para


def convert_md_to_docx(md_path: Path, out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)

    lines = md_path.read_text(encoding='utf-8').splitlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # Filter separator rows
        data_rows = [r for r in table_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
        if not data_rows:
            in_table = False
            table_rows = []
            return
        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip().strip('|').split('|')]
            parsed.append(cells)
        if not parsed:
            in_table = False
            table_rows = []
            return
        col_count = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=col_count)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        for ri, row in enumerate(parsed):
            for ci, cell_text in enumerate(row):
                if ci >= col_count:
                    break
                cell = tbl.cell(ri, ci)
                cell.text = ''
                para = cell.paragraphs[0]
                run = para.add_run(cell_text)
                run.font.size = Pt(9)
                if ri == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    tc_pr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), '0A6EBD')
                    tc_pr.append(shd)
        doc.add_paragraph()
        in_table = False
        table_rows = []

    def flush_code(lines_buf):
        if not lines_buf:
            return
        text = '\n'.join(lines_buf)
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
        p.paragraph_format.left_indent = Inches(0.3)
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'F4F6F8')
        pPr.append(shd)

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                flush_code(code_lines)
                code_lines = []
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^---+$', line.strip()):
            if in_table:
                flush_table()
            add_horizontal_rule(doc)
            i += 1
            continue

        # Table rows
        if line.strip().startswith('|'):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.*)', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown link syntax from headings
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            h = doc.add_heading(text, level=level)
            set_heading_style(h, level)
            if level == 1:
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(8)
                for run in h.runs:
                    run.font.size = Pt(18)
            elif level == 2:
                h.paragraph_format.space_before = Pt(14)
                h.paragraph_format.space_after = Pt(6)
                for run in h.runs:
                    run.font.size = Pt(14)
            elif level == 3:
                h.paragraph_format.space_before = Pt(10)
                h.paragraph_format.space_after = Pt(4)
                for run in h.runs:
                    run.font.size = Pt(12)
            else:
                for run in h.runs:
                    run.font.size = Pt(11)
            i += 1
            continue

        # Bullet list
        bullet_match = re.match(r'^(\s*)[-*+]\s+(.*)', line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text = bullet_match.group(2)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = add_paragraph_with_inline(doc, text)
            p.style = 'List Bullet' if indent == 0 else 'List Bullet 2'
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\s*)\d+\.\s+(.*)', line)
        if num_match:
            text = num_match.group(2)
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = add_paragraph_with_inline(doc, text)
            p.style = 'List Number'
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Blockquote
        if line.startswith('>'):
            text = line.lstrip('> ').strip()
            text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
            p = add_paragraph_with_inline(doc, text)
            p.paragraph_format.left_indent = Inches(0.4)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph
        text = line.strip()
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        if text:
            p = add_paragraph_with_inline(doc, text)
            p.paragraph_format.space_after = Pt(4)

        i += 1

    if in_table:
        flush_table()
    if in_code_block and code_lines:
        flush_code(code_lines)

    doc.save(str(out_path))
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    base = Path(__file__).parent.parent
    md_file = base / 'PROJECT_REPORT.md'
    out_file = base / 'ClaimIQ_Project_Report.docx'
    convert_md_to_docx(md_file, out_file)
